"""PDF to DOCX conversion — basic text and paragraph extraction."""

from io import BytesIO

import fitz
from docx import Document


def _normalize_block_text(text: str) -> str:
    """Collapse internal whitespace while preserving readable spacing."""
    return " ".join(text.split())


def pdf_to_docx(pdf_bytes: bytes) -> bytes:
    """
    Convert PDF bytes to DOCX bytes.

    Extracts positioned text blocks per page, sorted top-to-bottom then
    left-to-right. Each non-empty block becomes a paragraph; page breaks
    separate pages.

    @param pdf_bytes: Raw PDF file content
    @returns: Generated .docx file as bytes
    @raises ValueError: If the PDF cannot be opened or parsed
    """
    try:
        pdf = fitz.open(stream=pdf_bytes, filetype="pdf")
    except Exception as exc:
        raise ValueError("Could not read PDF file.") from exc

    doc = Document()
    has_content = False

    try:
        for page_index, page in enumerate(pdf):
            if page_index > 0:
                doc.add_page_break()

            blocks = page.get_text("blocks")
            # blocks: (x0, y0, x1, y1, text, block_no, block_type)
            text_blocks = [b for b in blocks if b[6] == 0 and b[4].strip()]
            text_blocks.sort(key=lambda b: (round(b[1], 1), round(b[0], 1)))

            for block in text_blocks:
                text = _normalize_block_text(block[4])
                if text:
                    doc.add_paragraph(text)
                    has_content = True
    finally:
        pdf.close()

    if not has_content:
        doc.add_paragraph(
            "No extractable text was found in this PDF. "
            "Scanned or image-only documents are not supported."
        )

    buffer = BytesIO()
    doc.save(buffer)
    return buffer.getvalue()
